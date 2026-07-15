from __future__ import annotations

from dataclasses import dataclass
import importlib.util
import os
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Optional

from presentation_agent.evaluation.adapters import evaluation_runtime_status
from presentation_agent.io import read_json, write_json


DEFAULT_HOME = Path.home() / "PresentationAgent"
DEFAULT_WORKSPACE = DEFAULT_HOME / "workspaces" / "default"


@dataclass(frozen=True)
class Workspace:
    root: Path

    @property
    def config_path(self) -> Path:
        return self.root / "config.toml"

    @property
    def data_dir(self) -> Path:
        return self.root / "data"

    @property
    def runs_dir(self) -> Path:
        return self.root / "runs"

    @property
    def artifacts_dir(self) -> Path:
        return self.root / "artifacts"

    @property
    def logs_dir(self) -> Path:
        return self.root / "logs"

    def run_dir(self, run: str) -> Path:
        candidate = Path(run).expanduser()
        if candidate.is_absolute() or candidate.exists() or "/" in run:
            return candidate.resolve()
        return (self.runs_dir / run).resolve()


def resolve_workspace(path: Optional[str | Path] = None, *, start: Optional[Path] = None) -> Workspace:
    if path:
        return Workspace(Path(path).expanduser().resolve())
    env = os.environ.get("PRESENTATION_AGENT_WORKSPACE")
    if env:
        return Workspace(Path(env).expanduser().resolve())
    found = find_workspace(start or Path.cwd())
    if found:
        return found
    return Workspace(DEFAULT_WORKSPACE.expanduser().resolve())


def find_workspace(start: Path) -> Optional[Workspace]:
    current = start.expanduser().resolve()
    if current.is_file():
        current = current.parent
    for candidate in [current, *current.parents]:
        marker = candidate / ".presentation-agent"
        if marker.exists() and marker.is_dir():
            return Workspace(marker.resolve())
    return None


def init_workspace(workspace: Workspace, repo_root: Path, *, force: bool = False) -> dict[str, object]:
    workspace.root.mkdir(parents=True, exist_ok=True)
    workspace.data_dir.mkdir(parents=True, exist_ok=True)
    workspace.runs_dir.mkdir(parents=True, exist_ok=True)
    workspace.artifacts_dir.mkdir(parents=True, exist_ok=True)
    workspace.logs_dir.mkdir(parents=True, exist_ok=True)

    config_created = _write_config(workspace, repo_root, force=force)
    global_created = _seed_global_state(workspace, repo_root, force=force)
    agent_count = _seed_agent_memory(workspace, repo_root, force=force)
    learning_created = _touch(workspace.data_dir / "learning" / "events.jsonl")

    return {
        "workspace": str(workspace.root),
        "config_created": config_created,
        "global_state_created": global_created,
        "agent_count": agent_count,
        "learning_events_created": learning_created,
    }


def workspace_status(workspace: Workspace, repo_root: Path) -> dict[str, object]:
    checks = []
    checks.append(_check("workspace_root", workspace.root.exists(), str(workspace.root)))
    checks.append(_check("config", workspace.config_path.exists(), str(workspace.config_path)))
    checks.append(_check("data", workspace.data_dir.exists(), str(workspace.data_dir)))
    checks.append(_check("runs", workspace.runs_dir.exists(), str(workspace.runs_dir)))
    checks.append(_check("artifacts", workspace.artifacts_dir.exists(), str(workspace.artifacts_dir)))
    checks.append(_check("repo_configs", (repo_root / "configs" / "agents.json").exists(), str(repo_root / "configs" / "agents.json")))
    checks.append(_check("repo_skills", (repo_root / "skills").exists(), str(repo_root / "skills")))
    production = _production_runtime_status()
    checks.extend(production["checks"])
    agents = _agent_ids(repo_root)
    missing_memory = [
        agent_id for agent_id in agents
        if not (workspace.data_dir / "agents" / agent_id / "memory.json").exists()
    ]
    checks.append(_check("agent_memory", not missing_memory, ", ".join(missing_memory) if missing_memory else "ok"))
    evaluation = evaluation_runtime_status(repo_root)
    ready_formats = [
        name
        for name, status in evaluation["formats"].items()
        if status["ready"]
    ]
    unavailable_dependencies = [
        item["name"]
        for item in evaluation["dependencies"]
        if item["status"] != "ok"
    ]
    evaluation_detail = (
        f"ready formats: {', '.join(ready_formats) or 'none'}; "
        f"unavailable: {', '.join(unavailable_dependencies) or 'none'}"
    )
    checks.append(
        _check(
            "evaluation_runtime",
            bool(evaluation["ok"]),
            evaluation_detail,
            required=False,
        )
    )
    return {
        "ok": all(
            item["status"] == "ok"
            for item in checks
            if item.get("required", True)
        ),
        "repo": str(repo_root),
        "workspace": str(workspace.root),
        "checks": checks,
        "production": production,
        "evaluation": evaluation,
    }


def _production_runtime_status() -> dict[str, object]:
    """Probe the real PDF/DOCX path used by Chinese report production."""

    pdfplumber_ready = importlib.util.find_spec("pdfplumber") is not None
    docx_ready = importlib.util.find_spec("docx") is not None
    bundled_soffice = (
        Path.home()
        / ".cache/codex-runtimes/codex-primary-runtime/dependencies/bin/soffice"
    )
    soffice = (
        os.environ.get("PRESENTATION_AGENT_SOFFICE")
        or shutil.which("soffice")
        or (str(bundled_soffice) if bundled_soffice.is_file() else None)
    )
    checks = [
        _check(
            "pdf_ingestion",
            pdfplumber_ready,
            "pdfplumber importable" if pdfplumber_ready else "pdfplumber is required to parse source PDFs",
        ),
        _check(
            "docx_generation",
            docx_ready,
            "python-docx importable" if docx_ready else "python-docx is required to generate reports",
        ),
    ]
    smoke_ok = False
    smoke_detail = "not run because python-docx, pdfplumber, or soffice is unavailable"
    if pdfplumber_ready and docx_ready and soffice:
        smoke_ok, smoke_detail = _chinese_document_smoke_test(Path(soffice))
    checks.append(_check("chinese_document_render", smoke_ok, smoke_detail))
    return {
        "ok": all(item["status"] == "ok" for item in checks),
        "checks": checks,
    }


def _chinese_document_smoke_test(soffice: Path) -> tuple[bool, str]:
    try:
        import pdfplumber
        from docx import Document

        with tempfile.TemporaryDirectory(prefix="presentation-agent-doctor-") as raw_dir:
            temp = Path(raw_dir)
            source = temp / "chinese-smoke.docx"
            document = Document()
            document.add_heading("中文标题", level=1)
            document.add_paragraph("战略分析中文渲染测试")
            document.save(source)
            completed = subprocess.run(
                [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(temp), str(source)],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
            pdf_path = temp / "chinese-smoke.pdf"
            if completed.returncode != 0 or not pdf_path.exists():
                detail = (completed.stderr or completed.stdout or "no PDF produced").strip()
                return False, f"LibreOffice DOCX→PDF smoke test failed: {detail}"
            with pdfplumber.open(pdf_path) as pdf:
                rendered_text = "\n".join((page.extract_text() or "") for page in pdf.pages)
            if "中文标题" not in rendered_text or "战略分析" not in rendered_text:
                return False, "DOCX→PDF completed but Chinese glyph text was missing"
            return True, "generated and rendered a Chinese DOCX successfully"
    except Exception as exc:
        return False, f"Chinese document smoke test failed: {exc}"


def _write_config(workspace: Workspace, repo_root: Path, *, force: bool) -> bool:
    if workspace.config_path.exists() and not force:
        return False
    content = "\n".join([
        "# Presentation Agent workspace",
        f"repo_root = {str(repo_root)!r}",
        "data_dir = 'data'",
        "runs_dir = 'runs'",
        "artifacts_dir = 'artifacts'",
        "",
    ])
    workspace.config_path.write_text(content, encoding="utf-8")
    return True


def _seed_global_state(workspace: Workspace, repo_root: Path, *, force: bool) -> bool:
    target = workspace.data_dir / "global" / "state.json"
    if target.exists() and not force:
        return False
    source = repo_root / "data" / "global" / "state.json"
    seed = read_json(source, default={})
    write_json(target, seed if isinstance(seed, dict) else {})
    return True


def _seed_agent_memory(workspace: Workspace, repo_root: Path, *, force: bool) -> int:
    count = 0
    for agent_id in _agent_ids(repo_root):
        agent_dir = workspace.data_dir / "agents" / agent_id
        agent_dir.mkdir(parents=True, exist_ok=True)
        memory_path = agent_dir / "memory.json"
        if force or not memory_path.exists():
            source = repo_root / "data" / "agents" / agent_id / "memory.json"
            seed = read_json(source, default={"items": []})
            write_json(memory_path, seed if isinstance(seed, dict) else {"items": []})
        _touch(agent_dir / "learning_log.jsonl")
        count += 1
    return count


def _touch(path: Path) -> bool:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return False
    path.write_text("", encoding="utf-8")
    return True


def _agent_ids(repo_root: Path) -> list[str]:
    config = read_json(repo_root / "configs" / "agents.json", default={})
    if not isinstance(config, dict):
        return []
    pipeline = config.get("pipeline", {})
    active = list(pipeline.get("stages", []))
    active.extend(pipeline.get("optional_workers", []))
    ids = ["manager"] if config.get("control_plane") else []
    ids.extend(str(agent_id) for agent_id in active if agent_id)
    if not ids:
        agents = config.get("agents", [])
        ids = [item["id"] for item in agents if isinstance(item, dict) and item.get("id")]
    return ids


def _check(
    name: str,
    ok: bool,
    detail: str,
    *,
    required: bool = True,
) -> dict[str, object]:
    result: dict[str, object] = {
        "name": name,
        "status": "ok" if ok else "missing",
        "detail": detail,
    }
    if not required:
        result["required"] = False
    return result
