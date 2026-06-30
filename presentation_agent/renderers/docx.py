"""Document render backend — formatted_material.v1 -> .docx via python-docx.

Mirrors the field conventions used by the PPT backend so agent4/agent5 emit one
shared `material_units` shape regardless of carrier:

    unit = {
        "unit_id": str,
        "headline": str,                # the McKinsey action title (a full insight sentence)
        "layout_or_structure": {"layout_type": "..."},
        "finalized_content": {
            "primary_text": str,
            "body": str,
            "supporting_points": [str | "标题：说明"],
            "section_label": str,
            "tables": [{"headers": [...], "rows": [[...]]}],
            "source": str,
        },
        "visual_object": {"visual_type": "donut|matrix_2x2|...", "data_fields": [...]},
        "source_display": {"footer": str},
    }

Each unit becomes a document section: a navy action-title heading, the body,
bulleted supporting points / bullet groups, quote blocks, formulas, a real Word
table when the unit carries tabular data, and a small source footer line.
Chart-bearing units can also produce a simple static PNG figure from
visual_object.chart_spec or data_fields so DOCX reports do not collapse into
pure text.

python-docx is an OPTIONAL dependency: if it is missing we return a skipped
result instead of crashing.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Optional

from presentation_agent.renderers.base import RenderResult

CAPABILITY_ID = "format.document"
REQUIRED_TOOLS = ("python-docx",)

# McKinsey palette (RGB tuples; python-docx wants hex strings / RGBColor)
_NAVY = (0x05, 0x1C, 0x2C)
_ACCENT_BLUE = (0x00, 0x6B, 0xA6)
_MED_GRAY = (0x66, 0x66, 0x66)
_DARK_GRAY = (0x33, 0x33, 0x33)
_LIGHT_FILL = "F2F6F9"  # table header fill
_PALETTE = ["#006BA6", "#007A53", "#D46A00", "#C62828", "#666666"]


# ---- shared content extraction (kept consistent with renderers/ppt.py) -------


def _content(unit: dict[str, Any]) -> dict[str, Any]:
    return unit.get("finalized_content") or {}


def _headline(unit: dict[str, Any]) -> str:
    c = _content(unit)
    return (unit.get("headline") or c.get("primary_text") or c.get("body") or "").strip()


def _supporting_points(content: dict[str, Any]) -> list[str]:
    pts = content.get("supporting_points") or []
    if isinstance(pts, str):
        pts = [pts]
    out = []
    for p in pts:
        if isinstance(p, (list, tuple)):
            out.append("：".join(str(x) for x in p))
        else:
            out.append(str(p))
    return [p for p in out if p.strip()]


def _bullet_groups(content: dict[str, Any]) -> list[dict[str, Any]]:
    groups = content.get("bullet_groups") or []
    return [g for g in groups if isinstance(g, dict)]


def _quote_blocks(content: dict[str, Any]) -> list[dict[str, Any]]:
    quotes = content.get("quote_blocks") or []
    return [q for q in quotes if isinstance(q, dict)]


def _formula_lines(content: dict[str, Any]) -> list[str]:
    lines = content.get("formula_lines") or []
    if isinstance(lines, str):
        lines = [lines]
    return [str(x) for x in lines if str(x).strip()]


def _source(unit: dict[str, Any], content: dict[str, Any]) -> str:
    sd = unit.get("source_display") or {}
    if isinstance(sd, dict) and sd.get("footer"):
        return str(sd["footer"])
    return content.get("source", "") if isinstance(content, dict) else ""


def _layout(unit: dict[str, Any]) -> str:
    los = unit.get("layout_or_structure") or {}
    return (los.get("layout_type") or unit.get("layout_type") or "").strip().lower()


def _table_from_unit(unit: dict[str, Any], content: dict[str, Any]) -> Optional[tuple[list[str], list[list[str]]]]:
    """Return (headers, rows) if the unit carries tabular data, else None.

    Sources, in priority order: finalized_content.tables, then a chart's
    visual_object.data_fields (chart degraded to a data table)."""
    tables = content.get("tables") or []
    if tables and isinstance(tables[0], dict):
        headers = [str(h) for h in (tables[0].get("headers") or ["项目", "说明"])]
        rows = [[str(cell) for cell in r] for r in (tables[0].get("rows") or [])]
        if rows:
            return headers, rows
    if tables and isinstance(tables[0], (list, tuple)):
        rows = [[str(cell) for cell in r] for r in tables]
        if rows:
            return ["项目", "数值", "说明"][: len(rows[0])] or ["项目"], rows

    visual = unit.get("visual_object") or {}
    vt = (visual.get("visual_type") or "").lower()
    fields = visual.get("data_fields") or []
    if vt in ("donut", "pie", "grouped_bar", "bar", "data_table", "table_insight") and fields:
        rows = []
        for item in fields:
            if isinstance(item, dict):
                label = str(item.get("label", item.get("name", "")))
                val = item.get("value", item.get("pct", item.get("percent", "")))
                sub = str(item.get("sub_label", item.get("sub", "")))
                rows.append([label, str(val), sub])
            elif isinstance(item, (list, tuple)):
                rows.append([str(x) for x in item][:3])
            else:
                rows.append([str(item), "", ""])
        if rows:
            width = max(len(r) for r in rows)
            headers = (["指标", "数值", "说明"])[:width]
            rows = [r + [""] * (width - len(r)) for r in rows]
            return headers, rows
    return None


# ---- low-level docx styling helpers ------------------------------------------


def _set_run_color(run, rgb: tuple[int, int, int]) -> None:
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor(*rgb)


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _add_action_title(doc, text: str, level: int = 1) -> None:
    from docx.shared import Pt

    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.bold = True
    _set_run_color(run, _NAVY)
    run.font.size = Pt(15 if level == 1 else 13)


def _add_body(doc, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    _set_run_color(run, _DARK_GRAY)


def _add_points(doc, points: list[str]) -> None:
    from docx.shared import Pt

    for p in points:
        para = doc.add_paragraph(style="List Bullet")
        # split "标题：说明" -> bold lead + normal tail
        lead, _, tail = p.partition("：")
        if tail:
            r1 = para.add_run(lead.strip())
            r1.bold = True
            _set_run_color(r1, _ACCENT_BLUE)
            r2 = para.add_run("：" + tail.strip())
            _set_run_color(r2, _DARK_GRAY)
        else:
            r = para.add_run(p)
            _set_run_color(r, _DARK_GRAY)


def _add_bullet_groups(doc, groups: list[dict[str, Any]]) -> None:
    from docx.shared import Pt

    for group in groups:
        title = str(group.get("group_title") or group.get("title") or "").strip()
        if title:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(10)
            _set_run_color(r, _ACCENT_BLUE)
        bullets = group.get("bullets") or group.get("items") or []
        if isinstance(bullets, str):
            bullets = [bullets]
        _add_points(doc, [str(b) for b in bullets if str(b).strip()])


def _add_formula_lines(doc, lines: list[str]) -> None:
    from docx.shared import Pt

    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(str(line))
        r.bold = True
        r.font.size = Pt(10)
        _set_run_color(r, _NAVY)


def _add_quote_blocks(doc, quotes: list[dict[str, Any]]) -> None:
    from docx.shared import Pt

    for q in quotes:
        text = str(q.get("quote") or q.get("text") or "").strip()
        if not text:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"“{text.strip('“”')}”")
        r.italic = True
        r.font.size = Pt(9)
        _set_run_color(r, _ACCENT_BLUE)
        attr = str(q.get("attribution") or q.get("source") or "").strip()
        if attr:
            ap = doc.add_paragraph()
            ar = ap.add_run(f"-- {attr}")
            ar.font.size = Pt(8)
            _set_run_color(ar, _MED_GRAY)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    from docx.shared import Pt

    ncol = max(len(headers), max((len(r) for r in rows), default=1))
    headers = (headers + [""] * ncol)[:ncol]
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        run = hdr[j].paragraphs[0].add_run(str(h))
        run.bold = True
        _set_run_color(run, _NAVY)
        run.font.size = Pt(10)
        _shade_cell(hdr[j], _LIGHT_FILL)
    for r in rows:
        cells = table.add_row().cells
        r = (list(r) + [""] * ncol)[:ncol]
        for j, val in enumerate(r):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            _set_run_color(run, _DARK_GRAY)
            run.font.size = Pt(10)


def _add_source(doc, source: str) -> None:
    from docx.shared import Pt

    if not source:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"来源：{source}")
    run.italic = True
    run.font.size = Pt(8)
    _set_run_color(run, _MED_GRAY)


def _font(size: int):
    from PIL import ImageFont

    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for path in candidates:
        try:
            if Path(path).exists():
                return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _to_float(v: Any) -> float:
    try:
        if isinstance(v, str):
            v = v.replace("%", "").replace(",", "").strip()
        return float(v)
    except Exception:
        return 0.0


def _short(text: Any, limit: int = 18) -> str:
    s = str(text)
    return s if len(s) <= limit else s[: limit - 1] + "…"


def _fallback_items(unit: dict[str, Any]) -> list[tuple[str, float]]:
    visual = unit.get("visual_object") or {}
    fields = visual.get("data_fields") or []
    out: list[tuple[str, float]] = []
    for item in fields:
        if isinstance(item, dict):
            out.append((str(item.get("label", item.get("name", ""))), _to_float(item.get("value", item.get("pct", 0)))))
        elif isinstance(item, (list, tuple)) and item:
            out.append((str(item[0]), _to_float(item[1] if len(item) > 1 else 0)))
    return out


def _draw_report_chart(unit: dict[str, Any], out_dir: Path) -> Optional[Path]:
    """Create a simple McKinsey-style PNG figure for DOCX embedding.

    This is intentionally conservative: it turns common chart_specs into a
    clean static figure so document outputs contain a real visual block even
    before a richer chart renderer exists.
    """
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None

    visual = unit.get("visual_object") or {}
    spec = visual.get("chart_spec") or {}
    layout = _layout(unit) or str(visual.get("visual_type") or "").lower()
    if not spec and not visual.get("data_fields"):
        return None

    w, h = 1400, 620
    margin_l, margin_r, margin_t, margin_b = 170, 70, 92, 95
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_font = _font(34)
    label_font = _font(22)
    small_font = _font(18)
    title = visual.get("title") or unit.get("headline") or ""
    d.text((45, 30), _short(title, 46), fill="#051C2C", font=title_font)

    chart_left, chart_top = margin_l, margin_t
    chart_w, chart_h = w - margin_l - margin_r, h - margin_t - margin_b
    d.line((chart_left, chart_top + chart_h, chart_left + chart_w, chart_top + chart_h), fill="#D7DEE5", width=2)

    def draw_hbars(items: list[tuple[str, float]]) -> None:
        items[:] = items[:8]
        max_v = max([v for _, v in items] + [1])
        gap = 18
        bar_h = max(24, min(48, (chart_h - gap * (len(items) - 1)) // max(len(items), 1)))
        y = chart_top + 20
        for idx, (name, val) in enumerate(items):
            color = _PALETTE[idx % len(_PALETTE)]
            d.text((40, y + 4), _short(name, 14), fill="#333333", font=small_font)
            bw = int(chart_w * (val / max_v))
            d.rounded_rectangle((chart_left, y, chart_left + bw, y + bar_h), radius=5, fill=color)
            d.text((chart_left + bw + 10, y + 5), f"{val:g}", fill="#333333", font=small_font)
            y += bar_h + gap

    if layout in ("horizontal_bar", "pareto", "donut", "pie") or "items" in spec or "segments" in spec:
        raw = spec.get("items") or spec.get("segments") or []
        items = [(str(x.get("name", x.get("label", ""))), _to_float(x.get("value", x.get("pct", 0)))) for x in raw if isinstance(x, dict)]
        draw_hbars(items or _fallback_items(unit))
    elif layout in ("line_chart", "timeline") and (spec.get("values") or spec.get("y_labels")):
        values = [_to_float(v) for v in (spec.get("values") or spec.get("y_labels") or [])]
        labels = spec.get("x_labels") or [str(i + 1) for i in range(len(values))]
        max_v, min_v = max(values + [1]), min(values + [0])
        span = max(max_v - min_v, 1)
        pts = []
        for i, val in enumerate(values):
            x = chart_left + int(chart_w * i / max(len(values) - 1, 1))
            y = chart_top + chart_h - int(chart_h * (val - min_v) / span)
            pts.append((x, y))
        if len(pts) >= 2:
            d.line(pts, fill="#006BA6", width=6)
        for (x, y), lab, val in zip(pts, labels, values):
            d.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#006BA6")
            d.text((x - 20, chart_top + chart_h + 14), _short(lab, 8), fill="#666666", font=small_font)
            d.text((x - 18, y - 34), f"{val:g}", fill="#333333", font=small_font)
    else:
        categories = spec.get("categories") or spec.get("periods") or []
        series = spec.get("series") or [spec.get("legend_label") or "value"]
        values = spec.get("values") or []
        if not categories and spec.get("panels"):
            panels = spec.get("panels")[:3]
            x0 = chart_left
            panel_w = chart_w // max(len(panels), 1)
            for pi, panel in enumerate(panels):
                vals = [_to_float(v) for v in (panel.get("values") or [])][:6]
                labels = panel.get("labels") or panel.get("categories") or [str(i + 1) for i in range(len(vals))]
                max_v = max(vals + [1])
                d.text((x0 + pi * panel_w, chart_top), _short(panel.get("title", ""), 15), fill="#051C2C", font=label_font)
                for i, val in enumerate(vals):
                    bw = max(18, panel_w // (len(vals) * 2 + 1))
                    x = x0 + pi * panel_w + 22 + i * (bw + 24)
                    bh = int((chart_h - 80) * val / max_v)
                    y = chart_top + chart_h - bh
                    d.rectangle((x, y, x + bw, chart_top + chart_h), fill=_PALETTE[i % len(_PALETTE)])
                    d.text((x - 8, chart_top + chart_h + 12), _short(labels[i], 5), fill="#666666", font=small_font)
        elif categories and values:
            n_cat, n_ser = len(categories), max(len(series), 1)
            max_v = max([_to_float(v) for row in values for v in (row if isinstance(row, list) else [row])] + [1])
            group_w = chart_w / max(n_cat, 1)
            bar_w = max(14, int(group_w / (n_ser + 1.3)))
            for i, cat in enumerate(categories[:8]):
                row = values[i] if i < len(values) and isinstance(values[i], list) else []
                x_base = chart_left + int(i * group_w) + 10
                for j in range(min(n_ser, 4)):
                    val = _to_float(row[j] if j < len(row) else 0)
                    bh = int((chart_h - 45) * val / max_v)
                    x = x_base + j * bar_w
                    y = chart_top + chart_h - bh
                    d.rectangle((x, y, x + bar_w - 3, chart_top + chart_h), fill=_PALETTE[j % len(_PALETTE)])
                d.text((x_base, chart_top + chart_h + 14), _short(cat, 7), fill="#666666", font=small_font)
        else:
            draw_hbars(_fallback_items(unit))

    caption = visual.get("caption") or visual.get("reader_takeaway") or ""
    source = visual.get("source_note") or (_content(unit).get("source_note") if isinstance(_content(unit), dict) else "") or _source(unit, _content(unit))
    if caption:
        d.text((45, h - 62), _short(caption, 80), fill="#333333", font=small_font)
    if source:
        d.text((45, h - 34), _short(f"来源：{source}", 90), fill="#666666", font=small_font)

    asset_dir = out_dir / "_docx_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    out = asset_dir / f"{unit.get('unit_id', 'chart')}.png"
    im.save(out)
    return out


def _add_chart_figure(doc, unit: dict[str, Any], content: dict[str, Any], out_dir: Path) -> None:
    from docx.shared import Inches, Pt

    img = _draw_report_chart(unit, out_dir)
    if not img:
        return
    doc.add_picture(str(img), width=Inches(6.7))
    visual = unit.get("visual_object") or {}
    caption = visual.get("caption") or content.get("figure_caption") or visual.get("reader_takeaway")
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(str(caption))
        r.font.size = Pt(8)
        _set_run_color(r, _MED_GRAY)


# ---- per-unit rendering ------------------------------------------------------


def _render_cover(doc, unit, content) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(_headline(unit) or "汇报标题")
    tr.bold = True
    tr.font.size = Pt(26)
    _set_run_color(tr, _NAVY)

    sub = content.get("body") or (_supporting_points(content)[:1] or [""])[0]
    if sub:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(str(sub))
        sr.font.size = Pt(13)
        _set_run_color(sr, _MED_GRAY)

    meta_bits = [b for b in (content.get("author", ""), content.get("date", "")) if b]
    if meta_bits:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run("  |  ".join(str(b) for b in meta_bits))
        mr.font.size = Pt(11)
        _set_run_color(mr, _MED_GRAY)
    doc.add_page_break()


def _render_divider(doc, unit, content) -> None:
    from docx.shared import Pt

    label = content.get("section_label", "")
    if label:
        lp = doc.add_paragraph()
        lr = lp.add_run(str(label))
        lr.bold = True
        lr.font.size = Pt(11)
        _set_run_color(lr, _ACCENT_BLUE)
    _add_action_title(doc, _headline(unit) or "章节", level=1)
    _add_body(doc, content.get("body", ""))


def _render_closing(doc, unit, content) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_headline(unit) or "谢谢")
    r.bold = True
    r.font.size = Pt(22)
    _set_run_color(r, _NAVY)
    msg = content.get("body") or content.get("primary_text") or ""
    if msg:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(str(msg))
        _set_run_color(mr, _MED_GRAY)


def _render_standard(doc, unit, content, draft: bool = False, out_dir: Optional[Path] = None) -> None:
    """Action-title section: heading + body + (table | bullets) + source.

    In draft fidelity we skip real Word tables (wireframe), keeping the numbers
    as plain bullets so agent4's output is fast and obviously low-fidelity."""
    _add_action_title(doc, _headline(unit) or "要点", level=2)
    body = content.get("primary_text") or content.get("body") or ""
    _add_body(doc, body)
    _add_formula_lines(doc, _formula_lines(content))

    table = _table_from_unit(unit, content) if not draft else None
    if table:
        _add_table(doc, table[0], table[1])
        # for table_insight, append the takeaways below the table
        insights = _supporting_points(content)[:3] if _layout(unit) == "table_insight" else []
        if insights:
            _add_points(doc, insights)
    else:
        groups = _bullet_groups(content)
        if groups:
            _add_bullet_groups(doc, groups)
        _add_points(doc, _supporting_points(content))

    if out_dir is not None:
        _add_chart_figure(doc, unit, content, out_dir)
    _add_quote_blocks(doc, _quote_blocks(content))

    _add_source(doc, _source(unit, content))


# ---- entrypoint --------------------------------------------------------------


def render_docx(
    material: dict[str, Any],
    out_dir: Path,
    fidelity: str = "final",
    file_stem: str = "deliverable",
) -> RenderResult:
    out_dir = Path(out_dir)
    units = material.get("material_units") or []
    if not units:
        return RenderResult(status="no_units", fmt="document", fidelity=fidelity)

    try:
        import docx  # python-docx  # noqa: F401
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:  # python-docx missing
        return RenderResult(
            status="skipped_missing_dep",
            fmt="document",
            fidelity=fidelity,
            unit_count=len(units),
            detail=f"需要 python-docx（pip install python-docx）：{exc}",
        )

    warnings: list[str] = []
    draft = fidelity == "draft"
    doc = Document()
    # base font
    try:
        normal = doc.styles["Normal"]
        normal.font.size = Pt(11)
        normal.font.name = "Calibri"
    except Exception:
        pass

    if draft:  # visible wireframe banner
        banner = doc.add_paragraph()
        br = banner.add_run("【草稿版 DRAFT — 待 format 精修，未含正式图表与排版】")
        br.bold = True
        _set_run_color(br, _MED_GRAY)

    total = len(units)
    for i, unit in enumerate(units):
        content = _content(unit)
        layout = _layout(unit)
        try:
            if layout == "cover" or (not layout and i == 0):
                _render_cover(doc, unit, content)
            elif layout == "section_divider":
                _render_divider(doc, unit, content)
            elif layout == "closing" or (not layout and i == total - 1 and total > 1):
                _render_closing(doc, unit, content)
            else:
                _render_standard(doc, unit, content, draft=draft, out_dir=out_dir)
        except Exception as exc:  # never let one unit kill the doc
            warnings.append(f"unit {unit.get('unit_id', i)} 降级为纯文本：{exc}")
            _add_action_title(doc, _headline(unit) or "要点", level=2)
            _add_body(doc, content.get("body", ""))

    suffix = "draft" if fidelity == "draft" else "final"
    out_path = out_dir / f"{file_stem}_{suffix}.docx"
    try:
        doc.save(str(out_path))
    except Exception as exc:
        return RenderResult(
            status="error", fmt="document", fidelity=fidelity, warnings=warnings, detail=str(exc)
        )

    size = out_path.stat().st_size if out_path.exists() else 0
    return RenderResult(
        status="rendered",
        fmt="document",
        fidelity=fidelity,
        output_path=str(out_path),
        file_bytes=size,
        unit_count=total,
        warnings=warnings,
    )
