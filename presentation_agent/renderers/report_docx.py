"""Semantic content-report renderer: report.v1 -> readable DOCX.

This renderer is intentionally independent from the legacy Format renderer.
It consumes report sections, narrative blocks, traceability registries and
appendices directly; it never interprets page/content/material units.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Iterable

from presentation_agent.renderers.base import RenderResult

CAPABILITY_ID = "report.content_docx"
PRESET_NAME = "standard_business_brief"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

_BLUE = "2E74B5"
_DARK_BLUE = "1F4D78"
_INK = "253746"
_MUTED = "666666"
_LIGHT_FILL = "F2F4F7"
_CALLOUT_FILL = "F4F6F9"
_CAVEAT_FILL = "FFF4CE"
_WHITE = "FFFFFF"


def _require_report_v1(report: dict[str, Any]) -> None:
    if report.get("agent_id") != "report" or report.get("schema") != "report.v1":
        raise ValueError("report_docx requires an artifact with agent_id=report and schema=report.v1")
    required = {
        "report_metadata",
        "executive_summary",
        "sections",
        "source_registry",
        "caveats_and_limits",
        "recommendations",
        "appendices",
        "format_handoff",
    }
    missing = sorted(required - report.keys())
    if missing:
        raise ValueError(f"report.v1 missing required fields: {', '.join(missing)}")


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for side, value in CELL_MARGIN_DXA.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _column_widths(columns: list[str], rows: list[list[Any]]) -> list[int]:
    count = max(1, len(columns))
    if count == 1:
        return [CONTENT_WIDTH_DXA]
    weights: list[int] = []
    for index in range(count):
        values = [columns[index]]
        values.extend(str(row[index]) for row in rows if index < len(row))
        weights.append(max(6, min(36, max((len(value) for value in values), default=6))))
    minimum = 1200 if count <= 4 else 800
    available = CONTENT_WIDTH_DXA - minimum * count
    if available < 0:
        minimum = CONTENT_WIDTH_DXA // count
        available = CONTENT_WIDTH_DXA - minimum * count
    total_weight = sum(weights)
    widths = [minimum + round(available * weight / total_weight) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_run_font(run: Any, *, size: float | None = None, color: str | None = None,
                  bold: bool | None = None, italic: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _style_document(doc: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(_INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, _BLUE, 16, 8),
        "Heading 2": (13, _BLUE, 12, 6),
        "Heading 3": (12, _DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    for name in ("Report Citation", "Report Note", "Report Callout"):
        if name not in doc.styles:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    citation = doc.styles["Report Citation"]
    citation.font.name = "Calibri"
    citation.font.size = Pt(8)
    citation.font.color.rgb = RGBColor.from_string(_MUTED)
    citation.font.italic = True
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    note = doc.styles["Report Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor.from_string(_INK)
    note.paragraph_format.space_after = Pt(6)
    callout = doc.styles["Report Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(11)
    callout.font.bold = True
    callout.font.color.rgb = RGBColor.from_string(_DARK_BLUE)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)


def _add_page_field(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run("第 ")
    _set_run_font(run, size=8, color=_MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    _set_run_font(run, size=8, color=_MUTED)


def _add_heading(doc: Any, text: str, level: int) -> Any:
    paragraph = doc.add_heading(level=level)
    paragraph.add_run(text)
    return paragraph


def _add_bullets(doc: Any, items: Iterable[Any]) -> None:
    for item in items:
        text = str(item).strip()
        if not text:
            continue
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(text)
        _set_run_font(run, color=_INK)


def _add_note_box(doc: Any, label: str, text: str, *, caveat: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, _CAVEAT_FILL if caveat else _CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.style = "Report Callout" if not caveat else "Report Note"
    lead = paragraph.add_run(f"{label}：")
    _set_run_font(lead, bold=True, color=_DARK_BLUE if not caveat else "7A5A00")
    body = paragraph.add_run(text)
    _set_run_font(body, color=_INK)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = 0


def _source_ids_for_block(report: dict[str, Any], block: dict[str, Any]) -> list[str]:
    # report.v1 keeps evidence refs on the block and the source registry at
    # report level; it intentionally has no second global claim-evidence map.
    return []


def _add_citation(doc: Any, evidence_refs: Iterable[str], source_ids: Iterable[str]) -> None:
    evidence = [str(item) for item in evidence_refs if str(item)]
    sources = [str(item) for item in source_ids if str(item)]
    if not evidence and not sources:
        return
    parts = []
    if evidence:
        parts.append("证据 " + "、".join(evidence))
    if sources:
        parts.append("来源 " + "、".join(f"[{item}]" for item in sources))
    paragraph = doc.add_paragraph(style="Report Citation")
    run = paragraph.add_run("；".join(parts))
    _set_run_font(run, size=8, color=_MUTED, italic=True)


def _add_data_table(doc: Any, title: str, columns: list[str], rows: list[list[Any]],
                    source_refs: Iterable[str] = (), notes: Iterable[str] = ()) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(title)
    _set_run_font(run, size=10, bold=True, color=_DARK_BLUE)
    count = max(1, len(columns))
    table = doc.add_table(rows=1, cols=count)
    table.style = "Table Grid"
    for index, heading in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, _LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(heading))
        _set_run_font(run, size=9.5, bold=True, color=_DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        normalized = list(row) + [""] * (count - len(row))
        for index, value in enumerate(normalized[:count]):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if count == 2 and index == 1 and len(str(value)) <= 12:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=9.5, color=_INK)
    _set_table_geometry(table, _column_widths(columns, rows))
    citation_bits = []
    refs = [str(item) for item in source_refs if str(item)]
    if refs:
        citation_bits.append("来源：" + "、".join(f"[{item}]" for item in refs))
    note_items = [str(item) for item in notes if str(item)]
    if note_items:
        citation_bits.append("注：" + "；".join(note_items))
    if citation_bits:
        paragraph = doc.add_paragraph(style="Report Citation")
        run = paragraph.add_run(" ".join(citation_bits))
        _set_run_font(run, size=8, color=_MUTED, italic=True)


def _render_narrative_block(doc: Any, report: dict[str, Any], block: dict[str, Any]) -> None:
    from docx.shared import Inches

    block_type = block.get("block_type")
    content = str(block.get("content") or "").strip()
    items = block.get("items") or []
    if block_type == "bullet_group":
        if content:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(content)
            _set_run_font(run, bold=True, color=_DARK_BLUE)
        _add_bullets(doc, items)
    elif block_type == "quote":
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3)
        run = paragraph.add_run(f"“{content.strip('“”')}”")
        _set_run_font(run, size=10.5, color=_DARK_BLUE, italic=True)
    elif block_type == "callout":
        _add_note_box(doc, "关键判断", content)
    elif block_type == "caveat":
        _add_note_box(doc, "边界", content, caveat=True)
    elif block_type == "method_note":
        _add_note_box(doc, "方法说明", content)
    elif block_type == "figure_placeholder":
        spec = block.get("figure_spec") or {}
        _render_figure_spec(doc, spec)
        if content:
            _add_note_box(doc, "图表说明", content)
    elif block_type == "table":
        table = block.get("table") or {}
        _add_data_table(
            doc,
            content or "表格",
            [str(value) for value in table.get("columns") or []],
            table.get("rows") or [],
            block.get("evidence_refs") or [],
        )
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        _set_run_font(run, color=_INK)
    _add_citation(
        doc,
        block.get("evidence_refs") or [],
        _source_ids_for_block(report, block),
    )


def _render_figure_spec(doc: Any, spec: dict[str, Any]) -> None:
    title = str(spec.get("title") or "图表")
    takeaway = str(spec.get("reader_takeaway") or "")
    _add_note_box(doc, "图表规格", f"{title}。{takeaway}")
    data = spec.get("data") or {}
    if data:
        paragraph = doc.add_paragraph(style="Report Citation")
        run = paragraph.add_run("数据：" + json.dumps(data, ensure_ascii=False, separators=(", ", ": ")))
        _set_run_font(run, size=8, color=_MUTED)
    caveats = spec.get("caveats") or []
    _add_citation(doc, caveats, spec.get("source_refs") or [])


def _render_executive_summary(doc: Any, summary: dict[str, Any]) -> None:
    _add_heading(doc, "执行摘要", 1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(summary.get("context") or ""))
    _set_run_font(run, color=_INK)
    _add_note_box(doc, "核心答案", str(summary.get("core_answer") or ""))
    _add_heading(doc, "关键发现", 2)
    _add_bullets(doc, summary.get("key_findings") or [])
    _add_heading(doc, "业务含义", 2)
    _add_bullets(doc, summary.get("implications") or [])
    _add_note_box(doc, "需要确认", str(summary.get("expected_action") or ""))


def _render_method_and_risks(doc: Any, report: dict[str, Any]) -> None:
    methodology = report.get("caveats_and_limits") or {}
    _add_heading(doc, "方法与边界", 1)
    for label, key in (("分析方法", "approach"),):
        paragraph = doc.add_paragraph()
        lead = paragraph.add_run(f"{label}：")
        _set_run_font(lead, bold=True, color=_DARK_BLUE)
        body = paragraph.add_run(str(methodology.get(key) or ""))
        _set_run_font(body, color=_INK)
    definitions = methodology.get("definitions") or []
    if definitions:
        _add_heading(doc, "关键定义", 2)
        _add_bullets(doc, definitions)
    limitations = methodology.get("limitations") or []
    if limitations:
        _add_heading(doc, "方法限制", 2)
        _add_bullets(doc, limitations)
    for heading, key in (("关键假设", "assumptions"), ("数据缺口", "data_gaps")):
        values = methodology.get(key) or []
        if values:
            _add_heading(doc, heading, 2)
            _add_bullets(doc, values)


def _render_recommendations(doc: Any, recommendations: list[dict[str, Any]]) -> None:
    if not recommendations:
        return
    _add_heading(doc, "建议", 1)
    for index, item in enumerate(recommendations, 1):
        paragraph = doc.add_paragraph()
        lead = paragraph.add_run(f"{index}. {item.get('statement', '')}")
        _set_run_font(lead, bold=True, color=_DARK_BLUE)
        rationale = str(item.get("rationale") or "")
        if rationale:
            body = paragraph.add_run(f" {rationale}")
            _set_run_font(body, color=_INK)
        conditions = item.get("conditions") or []
        if conditions:
            _add_bullets(doc, [f"成立条件：{condition}" for condition in conditions])


def _render_appendices(doc: Any, report: dict[str, Any]) -> None:
    appendices = report.get("appendices") or []
    sources = report.get("source_registry") or []
    trace_rows: list[list[str]] = []
    seen_claims: set[str] = set()
    for section in report.get("sections") or []:
        if not isinstance(section, dict):
            continue
        for block in section.get("narrative_blocks") or []:
            if not isinstance(block, dict):
                continue
            for claim_id in block.get("claim_ids") or []:
                claim_id = str(claim_id)
                if not claim_id or claim_id in seen_claims:
                    continue
                seen_claims.add(claim_id)
                trace_rows.append([
                    claim_id,
                    str(block.get("content") or ""),
                    "、".join(map(str, block.get("evidence_refs") or [])),
                    str(section.get("section_id") or ""),
                ])
    if not appendices and not sources and not trace_rows:
        return
    doc.add_page_break()
    _add_heading(doc, "附录", 1)
    for appendix in appendices:
        _add_heading(doc, str(appendix.get("title") or appendix.get("appendix_id") or "附录"), 2)
        _add_bullets(doc, [f"关联内容：{ref}" for ref in appendix.get("content_refs") or []])
    if trace_rows:
        _add_heading(doc, "主张与证据追溯", 2)
        _add_data_table(
            doc,
            "Claim traceability",
            ["Claim", "首次出现内容", "证据", "章节"],
            trace_rows,
        )
    if sources:
        _add_heading(doc, "来源清单", 2)
        rows = [
            [
                item.get("source_id", ""),
                item.get("citation", ""),
                item.get("locator", ""),
            ]
            for item in sources
        ]
        _add_data_table(doc, "Source registry", ["来源 ID", "来源", "定位"], rows)


def render_report_docx(
    report: dict[str, Any],
    out_dir: Path,
    *,
    file_stem: str = "report_content",
) -> RenderResult:
    """Render a report.v1 artifact into a basic, independently readable DOCX."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        _require_report_v1(report)
    except ValueError as exc:
        return RenderResult(status="error", fmt="document", fidelity="content", detail=str(exc))
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:
        return RenderResult(
            status="skipped_missing_dep",
            fmt="document",
            fidelity="content",
            unit_count=len(report.get("sections") or []),
            detail=f"需要 python-docx：{exc}",
        )

    doc = Document()
    _style_document(doc)
    metadata = report["report_metadata"]
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header_run = header.add_run(str(metadata.get("title") or "分析报告"))
    _set_run_font(header_run, size=8, color=_MUTED)
    _add_page_field(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(str(metadata.get("title") or "分析报告"))
    _set_run_font(title_run, size=24, color=_DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        f"{metadata.get('report_type', '')} | {metadata.get('audience', '')} | "
        f"版本 {metadata.get('version', '')}"
    )
    _set_run_font(subtitle_run, size=9.5, color=_MUTED)

    _render_executive_summary(doc, report["executive_summary"])
    for index, report_section in enumerate(report.get("sections") or [], 1):
        _add_heading(doc, f"{index}. {report_section.get('heading', '')}", 1)
        _add_note_box(doc, "本节判断", str(report_section.get("section_thesis") or ""))
        for block in report_section.get("narrative_blocks") or []:
            _render_narrative_block(doc, report, block)
        _add_note_box(doc, "本节结论", str(report_section.get("section_conclusion") or ""))
        transition = str(report_section.get("transition") or "")
        if transition:
            paragraph = doc.add_paragraph(style="Report Citation")
            run = paragraph.add_run("承接：" + transition)
            _set_run_font(run, size=8.5, color=_MUTED, italic=True)

    _render_method_and_risks(doc, report)
    _render_recommendations(doc, report.get("recommendations") or [])
    _render_appendices(doc, report)

    output_path = out_dir / f"{file_stem}.docx"
    try:
        doc.save(str(output_path))
    except Exception as exc:
        return RenderResult(
            status="error",
            fmt="document",
            fidelity="content",
            unit_count=len(report.get("sections") or []),
            detail=str(exc),
        )
    return RenderResult(
        status="rendered",
        fmt="document",
        fidelity="content",
        output_path=str(output_path),
        file_bytes=output_path.stat().st_size,
        unit_count=len(report.get("sections") or []),
        detail=f"preset={PRESET_NAME}",
    )
