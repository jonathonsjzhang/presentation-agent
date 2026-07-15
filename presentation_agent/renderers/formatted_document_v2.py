"""Independent formatted_material.v2(document) -> polished DOCX renderer."""

from __future__ import annotations

import json
import textwrap
from pathlib import Path
from typing import Any, Iterable

from presentation_agent.renderers.base import RenderResult
from presentation_agent.renderers.diagram import render_svg_with_png_fallback
from presentation_agent.renderers.report_docx import (
    _add_bullets,
    _add_data_table,
    _add_heading,
    _add_note_box,
    _add_page_field,
    _set_run_font,
    _style_document,
)

import re

CAPABILITY_ID = "format.document.v2"
PRESET_NAME = "standard_business_brief"


def _normalize_heading(heading: str) -> str:
    """Strip 'N. ' / 'N、' prefix so section_heading can match markdown ## lines."""
    return re.sub(r"^\d+[\.\、\s]+", "", heading).strip()


def _wrapped_title(text: Any, width: int = 42) -> str:
    return "\n".join(textwrap.wrap(str(text or ""), width=width)[:2])


def _validate(formatted: dict[str, Any], report: dict[str, Any]) -> None:
    if formatted.get("agent_id") != "format" or formatted.get("schema") != "formatted_material.v2":
        raise ValueError("formatted document renderer requires formatted_material.v2")
    if formatted.get("delivery_target") != "document":
        raise ValueError("formatted document renderer only accepts delivery_target=document")
    # Report authors the canonical manuscript, then Q&A appends the approved
    # question list while preserving the report.v1 schema. Format consumes
    # that enhanced manuscript, so both runtime producers are legitimate.
    if (
        report.get("agent_id") not in {"report", "qa_preparation"}
        or report.get("schema") != "report.v1"
    ):
        raise ValueError("formatted document renderer requires report.v1")
    if not str(report.get("report_markdown") or "").strip():
        raise ValueError("report.v1 requires report_markdown")


def _add_toc(doc: Any, report: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    _add_heading(doc, "目录", 1)
    entries = [
        line[3:].strip()
        for line in str(report.get("report_markdown") or "").splitlines()
        if line.startswith("## ")
    ]
    for index, entry in enumerate(entries):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(0 if index == 0 else 12)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(entry)
        _set_run_font(run, size=11 if index else 12, color="253746", bold=index == 0)


def _add_asset_trace(doc: Any, asset: dict[str, Any]) -> None:
    bits = [
        "Section: " + "、".join(asset.get("source_section_ids") or []),
        "Claim: " + "、".join(asset.get("source_claim_ids") or []) or "Claim: -",
        "Evidence: " + "、".join(asset.get("source_evidence_refs") or []) or "Evidence: -",
    ]
    paragraph = doc.add_paragraph(style="Report Citation")
    run = paragraph.add_run(" | ".join(bits))
    _set_run_font(run, size=8, color="666666", italic=True)
    note = str(asset.get("source_note") or "").strip()
    if note:
        paragraph = doc.add_paragraph(style="Report Citation")
        run = paragraph.add_run(note)
        _set_run_font(run, size=8, color="666666", italic=True)


def _chart_png(asset: dict[str, Any], path: Path) -> Path:
    from PIL import Image, ImageDraw

    from presentation_agent.renderers.diagram import _font

    data = _normalize_chart_data(asset)
    categories = list(data.get("categories") or [])
    series = data.get("series")
    chart_type = str(data.get("chart_type") or "bar").lower()
    if isinstance(series, list) and series and chart_type == "line":
        return _line_chart_png(asset, path, categories, series)
    if isinstance(series, list) and series:
        return _grouped_bar_chart_png(asset, path, categories, series)
    values = list(data.get("values") or [])
    if not categories or len(categories) != len(values):
        raise ValueError(f"{asset.get('asset_id')}: chart requires equally sized categories and values")
    numeric = [float(str(value).replace("%", "").replace(",", "")) for value in values]
    width, height = 1400, 700
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _wrapped_title(asset.get("title"), 40)
    draw.multiline_text((55, 24), title, fill="#0B2545", font=_font(32), spacing=8)
    maximum = max(numeric) if numeric else 1
    title_lines = max(1, title.count("\n") + 1)
    left, top, chart_width = 370, 115 + 42 * title_lines, 880
    bar_height, gap = 75, 55
    unit = str(data.get("unit") or "")
    value_labels = list(data.get("value_labels") or values)
    for index, (label, value, original) in enumerate(zip(categories, numeric, value_labels)):
        y = top + index * (bar_height + gap)
        draw.text((55, y + 20), str(label), fill="#253746", font=_font(24))
        bar_width = int(chart_width * value / maximum) if maximum else 0
        color = ("#006BA6", "#64748B", "#007A53", "#D46A00")[index % 4]
        draw.rounded_rectangle((left, y, left + bar_width, y + bar_height), radius=8, fill=color)
        value_label = f"{original}{unit}" if unit and not str(original).endswith(unit) else str(original)
        draw.text((left + bar_width + 18, y + 20), value_label, fill="#0B2545", font=_font(25))
    image.save(path, "PNG")
    return path


def _grouped_bar_chart_png(
    asset: dict[str, Any],
    path: Path,
    categories: list[Any],
    series: list[Any],
) -> Path:
    from PIL import Image, ImageDraw
    from presentation_agent.renderers.diagram import _font

    normalized = []
    for row in series:
        if not isinstance(row, dict):
            continue
        values = [_to_float(value) for value in row.get("values") or []]
        if len(values) == len(categories) and all(value is not None for value in values):
            normalized.append({"name": str(row.get("name") or ""), "values": values})
    if not categories or not normalized:
        raise ValueError(f"{asset.get('asset_id')}: grouped bar chart requires numeric series")

    width, height = 1400, 700
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _wrapped_title(asset.get("title"), 40)
    draw.multiline_text((55, 22), title, fill="#0B2545", font=_font(30), spacing=8)
    title_lines = max(1, title.count("\n") + 1)
    left, top, right, bottom = 115, 120 + 42 * title_lines, 1320, 565
    maximum = max(value for row in normalized for value in row["values"]) or 1
    colors = ("#006BA6", "#007A53", "#D46A00")
    group_width = (right - left) / len(categories)
    bar_width = min(58, max(18, int(group_width / (len(normalized) + 1.4))))
    for tick in range(5):
        value = maximum * tick / 4
        y = bottom - (bottom - top) * tick / 4
        draw.line((left, y, right, y), fill="#E2E8F0", width=1)
        draw.text((38, y - 10), _format_tick(value), fill="#64748B", font=_font(17))
    for series_index, row in enumerate(normalized):
        color = colors[series_index % len(colors)]
        legend_x = 105 + series_index * 260
        legend_y = top - 45
        draw.rectangle((legend_x, legend_y + 8, legend_x + 25, legend_y + 24), fill=color)
        draw.text((legend_x + 34, legend_y), row["name"], fill="#253746", font=_font(19))
        for category_index, value in enumerate(row["values"]):
            center = left + group_width * (category_index + 0.5)
            x = center + (series_index - (len(normalized) - 1) / 2) * bar_width - bar_width / 2
            y = bottom - (bottom - top) * value / maximum
            draw.rectangle((x, y, x + bar_width - 3, bottom), fill=color)
            draw.text((x, y - 24), _format_tick(value), fill="#253746", font=_font(15))
    for index, label in enumerate(categories):
        center = left + group_width * (index + 0.5)
        wrapped = "\n".join(textwrap.wrap(str(label), width=9)[:2])
        draw.multiline_text((center - 55, bottom + 16), wrapped, fill="#64748B", font=_font(16), spacing=2)
    image.save(path, "PNG")
    return path


def _numeric_range(value: Any) -> tuple[float | None, float | None]:
    numbers = re.findall(r"[-+]?\d+(?:\.\d+)?", str(value).replace(",", ""))
    parsed = [abs(float(number)) for number in numbers]
    if not parsed:
        return None, None
    return min(parsed), max(parsed)


def _normalize_chart_data(asset: dict[str, Any]) -> dict[str, Any]:
    """Project semantic worker payloads onto deterministic chart primitives."""
    data = dict(asset.get("data") or {})
    categories = data.get("categories")
    if isinstance(categories, list) and data.get("series"):
        return data
    if isinstance(categories, list) and isinstance(data.get("values"), list):
        return data
    if isinstance(categories, list):
        start = data.get("period_start_pct")
        end = data.get("period_end_pct")
        if isinstance(start, list) and isinstance(end, list):
            data["chart_type"] = "bar"
            data["series"] = [
                {"name": "期初", "values": start},
                {"name": "期末", "values": end},
            ]
            return data
    metrics = data.get("metrics")
    changes = data.get("change_range")
    if isinstance(metrics, list) and isinstance(changes, list) and len(metrics) == len(changes):
        highs = []
        for value in changes:
            _, high = _numeric_range(value)
            highs.append(high)
        if all(value is not None for value in highs):
            return {
                "chart_type": "bar",
                "categories": metrics,
                "values": highs,
                "value_labels": changes,
            }
    return data


def _line_chart_png(
    asset: dict[str, Any],
    path: Path,
    categories: list[Any],
    series: list[Any],
) -> Path:
    from PIL import Image, ImageDraw

    from presentation_agent.renderers.diagram import _font

    if not categories:
        raise ValueError(f"{asset.get('asset_id')}: line chart requires categories")
    normalized_series = []
    for row in series:
        if not isinstance(row, dict):
            continue
        values = [_to_float(value) for value in row.get("values") or []]
        if len(values) != len(categories):
            continue
        if any(value is not None for value in values):
            normalized_series.append(
                {"name": str(row.get("name") or ""), "values": values}
            )
    if not normalized_series:
        raise ValueError(f"{asset.get('asset_id')}: line chart requires numeric series")

    width, height = 1400, 700
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = _wrapped_title(asset.get("title"), 40)
    draw.multiline_text((55, 24), title, fill="#0B2545", font=_font(32), spacing=8)

    title_lines = max(1, title.count("\n") + 1)
    left, top, right, bottom = 105, 115 + 42 * title_lines, 1320, 565
    draw.line((left, bottom, right, bottom), fill="#94A3B8", width=2)
    draw.line((left, top, left, bottom), fill="#94A3B8", width=2)

    numeric = [
        value
        for row in normalized_series
        for value in row["values"]
        if value is not None
    ]
    minimum = min(numeric)
    maximum = max(numeric)
    if minimum == maximum:
        minimum -= 1
        maximum += 1
    span = maximum - minimum
    colors = ("#006BA6", "#007A53", "#D46A00")

    def xy(index: int, value: float) -> tuple[float, float]:
        x = left + (right - left) * index / max(1, len(categories) - 1)
        y = bottom - (bottom - top) * (value - minimum) / span
        return x, y

    for tick in range(5):
        value = minimum + span * tick / 4
        y = bottom - (bottom - top) * tick / 4
        draw.line((left, y, right, y), fill="#E2E8F0", width=1)
        draw.text((35, y - 12), _format_tick(value), fill="#64748B", font=_font(18))

    for series_index, row in enumerate(normalized_series):
        color = colors[series_index % len(colors)]
        previous: tuple[float, float] | None = None
        for index, value in enumerate(row["values"]):
            if value is None:
                previous = None
                continue
            point = xy(index, value)
            if previous is not None:
                draw.line((previous[0], previous[1], point[0], point[1]), fill=color, width=4)
            draw.ellipse((point[0] - 4, point[1] - 4, point[0] + 4, point[1] + 4), fill=color)
            previous = point
        legend_x = 125 + series_index * 205
        draw.rectangle((legend_x, 88, legend_x + 26, 104), fill=color)
        draw.text((legend_x + 34, 82), row["name"], fill="#253746", font=_font(20))

    label_step = max(1, len(categories) // 6)
    for index, label in enumerate(categories):
        if index not in (0, len(categories) - 1) and index % label_step:
            continue
        x = left + (right - left) * index / max(1, len(categories) - 1)
        draw.text((x - 45, bottom + 18), str(label)[:10], fill="#64748B", font=_font(16))

    note = str((asset.get("data") or {}).get("sampling_note") or "")
    if note:
        draw.text((55, 590), note, fill="#64748B", font=_font(18))
    image.save(path, "PNG")
    return path


def _to_float(value: Any) -> float | None:
    try:
        if value is None:
            return None
        return float(str(value).replace("%", "").replace(",", ""))
    except (TypeError, ValueError):
        return None


def _format_tick(value: float) -> str:
    if abs(value) >= 100:
        return f"{value:.0f}"
    return f"{value:.1f}".rstrip("0").rstrip(".")


def _matrix_png(asset: dict[str, Any], path: Path) -> Path:
    from PIL import Image, ImageDraw
    from presentation_agent.renderers.diagram import _font, _labels

    data = asset.get("data") or {}
    labels = _labels(data)
    dimensions = data.get("dimensions") or labels
    limitations = data.get("limitations") or []
    if not isinstance(dimensions, list) or len(dimensions) != 4 or any(
        not str(value).strip() for value in dimensions
    ):
        raise ValueError(
            f"{asset.get('asset_id')}: matrix requires exactly four non-empty dimensions/labels"
        )
    if limitations and (
        not isinstance(limitations, list) or len(limitations) != 4
    ):
        raise ValueError(
            f"{asset.get('asset_id')}: matrix limitations must contain exactly four items"
        )
    image = Image.new("RGB", (1200, 780), "white")
    draw = ImageDraw.Draw(image)
    title = _wrapped_title(asset.get("title"), 34)
    draw.multiline_text((45, 20), title, fill="#0B2545", font=_font(30), spacing=8)
    title_lines = max(1, title.count("\n") + 1)
    box_top = 105 + 40 * title_lines
    box_bottom = box_top + 500
    mid_y = box_top + 250
    draw.rectangle((170, box_top, 1090, box_bottom), outline="#64748B", width=3)
    draw.line((630, box_top, 630, box_bottom), fill="#CBD5E1", width=3)
    draw.line((170, mid_y, 1090, mid_y), fill="#CBD5E1", width=3)
    positions = (
        (205, box_top + 35),
        (665, box_top + 35),
        (205, mid_y + 35),
        (665, mid_y + 35),
    )
    for index, (label, position) in enumerate(zip(dimensions, positions)):
        draw.text(position, str(label), fill="#0B2545", font=_font(22))
        if index < len(limitations):
            wrapped = "\n".join(textwrap.wrap(str(limitations[index]), width=24)[:5])
            draw.multiline_text(
                (position[0], position[1] + 46),
                wrapped,
                fill="#475569",
                font=_font(17),
                spacing=6,
            )
    for key, xy in (("x_label", (500, box_bottom + 25)), ("y_label", (25, mid_y))):
        if data.get(key):
            draw.text(xy, str(data[key]), fill="#64748B", font=_font(20))
    image.save(path, "PNG")
    return path


def _render_visual_asset(
    doc: Any, asset: dict[str, Any], asset_dir: Path, *, strict: bool = False
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    asset_type = str(asset.get("asset_type") or "")
    title = str(asset.get("title") or asset.get("asset_id") or "视觉资产")
    data = asset.get("data") or {}
    image_path = Path(str(data.get("image_path") or ""))
    if image_path.is_file() and image_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.35))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(title)
        _set_run_font(run, size=9.5, color="1F4D78", bold=True)
    elif asset_type == "table":
        columns = [str(value) for value in data.get("columns") or data.get("headers") or []]
        rows = data.get("rows") or []
        if not columns:
            raise ValueError(f"{asset.get('asset_id')}: table requires columns/headers")
        _add_data_table(doc, title, columns, rows, asset.get("source_evidence_refs") or [], asset.get("caveats") or [])
    elif asset_type == "callout":
        _add_note_box(doc, title, str(asset.get("reader_takeaway") or ""))
    elif asset_type == "chart" and not _chart_data_ready(asset):
        if strict:
            raise ValueError(f"{asset.get('asset_id')}: chart data is not renderer-ready")
        _add_note_box(
            doc,
            title,
            str(asset.get("reader_takeaway") or "图表数据尚未提供，暂以要点呈现。"),
        )
    else:
        asset_dir.mkdir(parents=True, exist_ok=True)
        try:
            if asset_type == "chart":
                png_path = _chart_png(asset, asset_dir / f"{asset['asset_id']}.png")
            elif asset_type == "matrix":
                _, png_path = render_svg_with_png_fallback(asset, asset_dir)
                _matrix_png(asset, png_path)
            else:
                _, png_path = render_svg_with_png_fallback(asset, asset_dir)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(png_path), width=Inches(6.35))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"{title} — {asset.get('reader_takeaway', '')}")
            _set_run_font(run, size=9.5, color="1F4D78", bold=True)
        except Exception:
            if strict:
                raise
            _add_note_box(
                doc,
                title,
                str(asset.get("reader_takeaway") or f"「{title}」渲染失败，数据格式不兼容，暂以文本呈现。"),
            )
    _add_asset_trace(doc, asset)
    for caveat in asset.get("caveats") or []:
        _add_note_box(doc, "图表边界", str(caveat), caveat=True)


def _chart_data_ready(asset: dict[str, Any]) -> bool:
    data = _normalize_chart_data(asset)
    if not isinstance(data, dict):
        return False
    categories = data.get("categories")
    series = data.get("series")
    if isinstance(categories, list) and categories and isinstance(series, list) and series:
        return True
    values = data.get("values")
    return (
        isinstance(categories, list)
        and isinstance(values, list)
        and bool(categories)
        and len(categories) == len(values)
    )


def _assets_by_section(formatted: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    if "visuals" in formatted:
        result: dict[str, list[dict[str, Any]]] = {}
        for index, visual in enumerate(formatted.get("visuals") or [], 1):
            if not isinstance(visual, dict):
                continue
            data = visual.get("data") if isinstance(visual.get("data"), dict) else {}
            prepared = {
                "asset_id": str(visual.get("visual_evidence_id") or f"VIS-{index:02d}"),
                "asset_type": visual.get("type"),
                "title": visual.get("title"),
                "reader_takeaway": str(
                    data.get("text") or data.get("quote") or visual.get("title") or ""
                ),
                "data": data,
                "source_evidence_refs": visual.get("source_refs", []),
                "source_section_ids": [visual.get("section_heading")],
                "source_claim_ids": [],
            }
            key = _normalize_heading(str(visual.get("section_heading") or ""))
            result.setdefault(key, []).append(prepared)
        return result
    ordered_ids = formatted.get("render_plan", {}).get("asset_order") or []
    lookup = {item.get("asset_id"): item for item in formatted.get("visual_assets") or []}
    assets = [lookup[item] for item in ordered_ids if item in lookup]
    assets.extend(item for item in formatted.get("visual_assets") or [] if item not in assets)
    result: dict[str, list[dict[str, Any]]] = {}
    for asset in assets:
        for section_id in asset.get("source_section_ids") or []:
            result.setdefault(section_id, []).append(asset)
    return result


def _protected_caveats_for_section(
    formatted: dict[str, Any],
    report_section: dict[str, Any],
) -> list[str]:
    section_id = str(report_section.get("section_id") or "")
    unit_ids = {
        str(unit.get("unit_id"))
        for unit in formatted.get("delivery_units") or []
        if isinstance(unit, dict)
        and section_id in set(map(str, unit.get("source_section_ids") or []))
    }
    caveats: list[str] = []
    for row in formatted.get("caveat_preservation") or []:
        if not isinstance(row, dict):
            continue
        destinations = set(map(str, row.get("destination_unit_ids") or []))
        caveat = str(row.get("source_caveat") or "")
        if caveat and destinations & unit_ids and caveat not in caveats:
            caveats.append(caveat)
    return caveats


def _markdown_sections(markdown: str) -> list[tuple[str, list[str]]]:
    """Split canonical Markdown into H2 sections while preserving body lines."""

    sections: list[tuple[str, list[str]]] = []
    heading = ""
    body: list[str] = []
    for raw in markdown.splitlines():
        if raw.startswith("# ") and not raw.startswith("## "):
            continue
        if raw.startswith("## "):
            if heading or any(line.strip() for line in body):
                sections.append((heading, body))
            heading = raw[3:].strip()
            body = []
        else:
            body.append(raw)
    if heading or any(line.strip() for line in body):
        sections.append((heading, body))
    return sections


def _render_markdown_body(
    doc: Any,
    lines: list[str],
    marker_assets: dict[str, dict[str, Any]] | None = None,
    asset_dir: Path | None = None,
    *,
    strict: bool = False,
) -> set[str]:
    """Render a conservative Markdown subset without rewriting its wording."""

    marker_assets = marker_assets or {}
    rendered_assets: set[str] = set()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        marker_match = re.fullmatch(r"\[可视化论据：([A-Za-z0-9_-]+)\]", line.strip())
        if marker_match:
            evidence_id = marker_match.group(1)
            asset = marker_assets.get(evidence_id)
            if asset is not None and asset_dir is not None:
                _render_visual_asset(doc, asset, asset_dir, strict=strict)
                rendered_assets.add(str(asset.get("asset_id") or evidence_id))
            index += 1
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 2)
            index += 1
            continue
        if line.startswith("> "):
            _add_note_box(doc, "边界说明", line[2:].strip(), caveat=True)
            index += 1
            continue
        if line.startswith(("- ", "* ")):
            items: list[str] = []
            while index < len(lines) and lines[index].startswith(("- ", "* ")):
                items.append(lines[index][2:].strip())
                index += 1
            _add_bullets(doc, items)
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [
                [cell.strip() for cell in row.strip().strip("|").split("|")]
                for row in table_lines
            ]
            if len(rows) >= 2:
                columns = rows[0]
                data_rows = rows[2:] if set("".join(rows[1])) <= {"-", ":", " "} else rows[1:]
                _add_data_table(doc, "", columns, data_rows)
            continue
        paragraph_lines = [line]
        index += 1
        while (
            index < len(lines)
            and lines[index].strip()
            and not lines[index].startswith(("### ", "> ", "- ", "* ", "|", "[可视化论据："))
        ):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(" ".join(paragraph_lines))
        _set_run_font(run, color="253746")
    return rendered_assets


def _set_update_fields(doc: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _normalize_east_asia_font(doc: Any) -> None:
    """Enforce KaiTi for CJK and Arial for Latin text on every run."""
    from docx.oxml.ns import qn

    for style in doc.styles:
        if getattr(style, "_element", None) is not None and style._element.rPr is not None:
            fonts = style._element.rPr.get_or_add_rFonts()
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            fonts.set(qn("w:eastAsia"), "Kaiti SC")
    for paragraph in list(doc.paragraphs):
        for run in paragraph.runs:
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            fonts.set(qn("w:eastAsia"), "Kaiti SC")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                        fonts.set(qn("w:ascii"), "Arial")
                        fonts.set(qn("w:hAnsi"), "Arial")
                        fonts.set(qn("w:eastAsia"), "Kaiti SC")


def render_formatted_document_v2(
    formatted: dict[str, Any],
    report: dict[str, Any],
    out_dir: Path,
    *,
    file_stem: str = "report_formatted",
) -> RenderResult:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        _validate(formatted, report)
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:
        return RenderResult(status="error", fmt="document", fidelity="formatted", detail=str(exc))

    try:
        doc = Document()
        _style_document(doc)
        _set_update_fields(doc)
        metadata = report.get("report_metadata") or {}
        markdown = str(report.get("report_markdown") or "")
        markdown_title = next(
            (
                line[2:].strip()
                for line in markdown.splitlines()
                if line.startswith("# ") and not line.startswith("## ")
            ),
            "分析报告",
        )
        title = str(metadata.get("title") or markdown_title)
        section = doc.sections[0]
        _add_page_field(section.footer.paragraphs[0])

        # Compact opening: a single report title followed immediately by ES.
        title_paragraph = doc.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(10)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_paragraph.add_run(title)
        _set_run_font(run, size=24, color="1A1A1A", bold=True)

        section_assets = _assets_by_section(formatted)
        marker_assets = {
            str(asset.get("asset_id")): asset
            for assets in section_assets.values()
            for asset in assets
            if asset.get("asset_id")
        }
        # Build heading → section_id map from report.v1 sections so both
        # Path A (visuals keyed by section_heading) and Path B (visual_assets
        # keyed by source_section_ids) resolve correctly.
        section_id_map: dict[str, str] = {}
        for sec in report.get("sections") or []:
            sid = str(sec.get("section_id") or "")
            h = str(sec.get("heading") or "")
            if sid and h:
                section_id_map[_normalize_heading(h)] = sid
        markdown_sections = _markdown_sections(markdown)
        strict_render = bool(formatted.get("strict_render"))
        render_warnings: list[str] = []
        for section_index, (heading, body) in enumerate(markdown_sections):
            heading_paragraph = _add_heading(doc, heading, 1)
            if section_index == 0:
                heading_paragraph.paragraph_format.space_before = Pt(0)
            rendered_assets = _render_markdown_body(
                doc,
                body,
                marker_assets,
                out_dir / f"{file_stem}_assets",
                strict=strict_render,
            )
            norm = _normalize_heading(heading)
            sid = section_id_map.get(norm, "")
            # Path A: visuals keyed by normalized heading
            # Path B: visual_assets keyed by section_id
            for asset in (section_assets.get(norm) or []) + (section_assets.get(sid) or []):
                if str(asset.get("asset_id") or "") in rendered_assets:
                    continue
                try:
                    _render_visual_asset(
                        doc,
                        asset,
                        out_dir / f"{file_stem}_assets",
                        strict=strict_render,
                    )
                except Exception as exc:
                    render_warnings.append(
                        f"{asset.get('asset_id', '?')} ({asset.get('asset_type', '?')}): {exc}"
                    )
        _normalize_east_asia_font(doc)
        output_path = out_dir / f"{file_stem}.docx"
        doc.save(output_path)
        return RenderResult(
            status="rendered",
            fmt="document",
            fidelity="formatted",
            output_path=str(output_path),
            file_bytes=output_path.stat().st_size,
            unit_count=len(markdown_sections),
            warnings=render_warnings,
            detail=f"preset={PRESET_NAME}; visuals={len(formatted.get('visuals') or [])}"
            + (f"; failed={len(render_warnings)}" if render_warnings else ""),
        )
    except Exception as exc:
        return RenderResult(
            status="error",
            fmt="document",
            fidelity="formatted",
            unit_count=len(_markdown_sections(str(report.get("report_markdown") or ""))),
            detail=str(exc),
        )
