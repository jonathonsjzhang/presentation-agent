from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from presentation_agent.renderers.report_docx import (
    CELL_MARGIN_DXA,
    CONTENT_WIDTH_DXA,
    PRESET_NAME,
    TABLE_INDENT_DXA,
    render_report_docx,
)


ROOT = Path(__file__).resolve().parents[1]
FIXTURE = ROOT / "tests" / "fixtures" / "v0_3" / "report.v1.valid.json"
DOC_SKILL = Path(
    "/Users/zhangsijing/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.630.12135/skills/documents"
)


def read_report() -> dict:
    return json.loads(FIXTURE.read_text(encoding="utf-8"))


class ReportDocxTests(unittest.TestCase):
    def test_renderer_rejects_non_report_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_report_docx(
                {"agent_id": "format", "schema": "formatted_material.v2"},
                Path(temp_dir),
            )
        self.assertEqual(result.status, "error")
        self.assertIn("report.v1", result.detail)

    def test_report_fixture_generates_openable_extractable_docx(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_report_docx(read_report(), Path(temp_dir), file_stem="smoke")
            self.assertEqual(result.status, "rendered", result.detail)
            output = Path(result.output_path or "")
            self.assertTrue(output.is_file())
            self.assertGreater(output.stat().st_size, 10_000)
            with zipfile.ZipFile(output) as package:
                self.assertIn("word/document.xml", package.namelist())
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
            extracted = text + "\n" + table_text
            for expected in (
                "AI 助手用户留存改善机会",
                "执行摘要",
                "成果保存与回访共同出现",
                "方法与边界",
                "主张与证据追溯",
                "来源清单",
                "34%",
                "非随机分组",
            ):
                self.assertIn(expected, extracted)

    def test_preset_geometry_styles_and_tables_are_explicit(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_report_docx(read_report(), Path(temp_dir))
            document = Document(result.output_path)
            section = document.sections[0]
            self.assertEqual(PRESET_NAME, "standard_business_brief")
            self.assertAlmostEqual(section.page_width.inches, 8.5, places=2)
            self.assertAlmostEqual(section.page_height.inches, 11, places=2)
            self.assertAlmostEqual(section.left_margin.inches, 1, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Calibri")
            self.assertAlmostEqual(document.styles["Normal"].font.size.pt, 11, places=1)
            self.assertAlmostEqual(document.styles["Heading 1"].font.size.pt, 16, places=1)
            self.assertTrue(document.tables)
            for table in document.tables:
                tbl_pr = table._tbl.tblPr
                self.assertEqual(int(tbl_pr.find(qn("w:tblW")).get(qn("w:w"))), CONTENT_WIDTH_DXA)
                self.assertEqual(int(tbl_pr.find(qn("w:tblInd")).get(qn("w:w"))), TABLE_INDENT_DXA)
                grid_widths = [
                    int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.findall(qn("w:gridCol"))
                ]
                self.assertEqual(sum(grid_widths), CONTENT_WIDTH_DXA)
                for row in table.rows:
                    self.assertEqual(
                        [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells],
                        grid_widths,
                    )
                    for cell in row.cells:
                        tc_mar = cell._tc.tcPr.find(qn("w:tcMar"))
                        self.assertIsNotNone(tc_mar)
                        for side, value in CELL_MARGIN_DXA.items():
                            self.assertEqual(int(tc_mar.find(qn(f"w:{side}")).get(qn("w:w"))), value)

    @unittest.skipUnless(
        (DOC_SKILL / "render_docx.py").is_file()
        and shutil.which("soffice")
        and shutil.which("pdftoppm"),
        "LibreOffice/Poppler render dependencies are unavailable",
    )
    def test_generated_docx_renders_to_page_pngs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            result = render_report_docx(read_report(), temp, file_stem="visual")
            render_dir = temp / "rendered"
            env = dict(os.environ)
            env["TMPDIR"] = "/private/tmp"
            completed = subprocess.run(
                [
                    sys.executable,
                    str(DOC_SKILL / "render_docx.py"),
                    str(result.output_path),
                    "--output_dir",
                    str(render_dir),
                ],
                capture_output=True,
                text=True,
                env=env,
                timeout=90,
            )
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            pages = sorted(render_dir.glob("page-*.png"))
            self.assertGreaterEqual(len(pages), 3)
            self.assertTrue(all(page.stat().st_size > 10_000 for page in pages))


if __name__ == "__main__":
    unittest.main()
