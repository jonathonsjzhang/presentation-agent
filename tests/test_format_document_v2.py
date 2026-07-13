from __future__ import annotations

import copy
import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from presentation_agent.io import read_json, write_json
from presentation_agent.renderers.base import render_material
from presentation_agent.renderers.base import RenderResult
from presentation_agent.renderers.formatted_document_v2 import render_formatted_document_v2
from presentation_agent.step import StepRunner

ROOT = Path(__file__).resolve().parents[1]
FORMATTED = ROOT / "tests/fixtures/v0_3/formatted_material.v2.valid.json"
REPORT = ROOT / "tests/fixtures/v0_3/report.v1.valid.json"


def load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


class FormattedDocumentV2Tests(unittest.TestCase):
    def test_dispatcher_renders_all_v2_targets(self) -> None:
        report = load(REPORT)
        for target, suffix in (
            ("document", ".docx"),
            ("html", ".html"),
            ("ppt", ".pptx"),
        ):
            formatted = load(FORMATTED)
            formatted["delivery_target"] = target
            with self.subTest(target=target), tempfile.TemporaryDirectory() as temp_dir:
                result = render_material(
                    formatted,
                    Path(temp_dir),
                    source_report=report,
                    expected_format=target,
                )
                self.assertEqual(result.status, "rendered", result.detail)
                self.assertTrue(str(result.output_path).endswith(suffix))
                self.assertTrue(Path(str(result.output_path)).is_file())

    def test_fixture_renders_polished_traceable_docx(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(load(FORMATTED), load(REPORT), Path(temp_dir))
            self.assertEqual(result.status, "rendered", result.detail)
            output = Path(result.output_path or "")
            self.assertTrue(output.is_file())
            with zipfile.ZipFile(output) as package:
                self.assertIn("word/document.xml", package.namelist())
                xml = package.read("word/document.xml").decode("utf-8")
                self.assertNotIn(">目录<", xml)
                self.assertIn("方法与边界", xml)
                self.assertIn('w:ascii="Arial"', xml)
                self.assertIn('w:eastAsia="Kaiti SC"', xml)
                footer_xml = "".join(
                    package.read(name).decode("utf-8")
                    for name in package.namelist()
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
                self.assertIn("PAGE", footer_xml)
            document = Document(output)
            extracted = "\n".join(p.text for p in document.paragraphs)
            extracted += "\n" + "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
            for text in (
                "AI 助手用户留存改善机会",
                "Executive Summary",
                "34%",
                "Section: 一、成果保存与回访共同出现，但因果仍待验证",
                "Evidence: E-Q-01",
                "结论与战略含义",
            ):
                self.assertIn(text, extracted)
            self.assertTrue(document.inline_shapes)
            self.assertEqual(document.paragraphs[0].text, "AI 助手用户留存改善机会")
            self.assertEqual(document.paragraphs[0].runs[0].font.size.pt, 24)
            self.assertEqual(document.paragraphs[1].text, "Executive Summary")

    def test_renderer_does_not_require_worker_generated_caveat_register(self) -> None:
        formatted = load(FORMATTED)
        formatted.pop("caveat_preservation")
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(formatted, load(REPORT), Path(temp_dir))
        self.assertEqual(result.status, "rendered", result.detail)

    def test_renderer_accepts_qa_enhanced_report_v1(self) -> None:
        report = load(REPORT)
        report["agent_id"] = "qa_preparation"
        report["report_markdown"] += (
            "\n## 听众可能追问的问题\n\n"
            "- 哪些反例会推翻当前判断？\n"
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(
                load(FORMATTED), report, Path(temp_dir)
            )
            self.assertEqual(result.status, "rendered", result.detail)
            self.assertTrue(Path(str(result.output_path)).is_file())

    def test_opening_visual_marker_is_replaced_in_place(self) -> None:
        from docx import Document

        report = load(REPORT)
        report["report_markdown"] = (
            "# 用户时长分析\n\n## Executive Summary\n\n"
            "历史用户时长持续上升。\n\n[可视化论据：VE-OPEN]\n\n"
            "## 一、增长判断\n\n变化并非单点波动。\n"
        )
        report["visual_evidence_placements"] = [
            {
                "id": "VE-OPEN",
                "claim": "历史用户时长持续上升",
                "purpose": "展示完整历史变化",
                "evidence_refs": ["E-01"],
                "data_type": "time_series",
                "required": True,
                "placement": "opening",
                "section_heading": "Executive Summary",
                "marker": "[可视化论据：VE-OPEN]",
            }
        ]
        formatted = load(FORMATTED)
        formatted["visuals"] = [
            {
                "visual_evidence_id": "VE-OPEN",
                "section_heading": "Executive Summary",
                "type": "chart",
                "title": "历史用户时长变化",
                "source_refs": ["E-01"],
                "required": True,
                "placement": "opening",
                "data": {
                    "chart_type": "line",
                    "categories": ["1月", "2月", "3月"],
                    "series": [{"name": "时长", "values": [8.3, 10.7, 15.0]}],
                },
            }
        ]
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(formatted, report, Path(temp_dir))
            self.assertEqual(result.status, "rendered", result.detail)
            document = Document(result.output_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            self.assertNotIn("[可视化论据：VE-OPEN]", "\n".join(paragraphs))
            self.assertLess(
                next(i for i, text in enumerate(paragraphs) if "历史用户时长变化" in text),
                paragraphs.index("一、增长判断"),
            )

    def test_empty_chart_falls_back_to_callout(self) -> None:
        formatted = load(FORMATTED)
        formatted["visuals"][0]["type"] = "chart"
        formatted["visuals"][0]["data"] = {
            "categories": [],
            "values": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(
                formatted,
                load(REPORT),
                Path(temp_dir),
            )
        self.assertEqual(result.status, "rendered", result.detail)

    def test_matrix_and_callout_use_supplied_data_and_png_fallback(self) -> None:
        formatted = load(FORMATTED)
        extras = [
            {
                "section_heading": "二、可复用成果提供了值得优先验证的回访理由",
                "type": "matrix",
                "title": "实验优先级矩阵",
                "data": {"items": ["成果复用", "提醒触达"], "x_label": "证据强度", "y_label": "业务价值"},
                "source_refs": ["E-Q-01", "E-I-01"]
            },
            {
                "section_heading": "二、可复用成果提供了值得优先验证的回访理由",
                "type": "callout",
                "title": "实验原则",
                "data": {},
                "source_refs": ["E-Q-01"]
            },
        ]
        formatted["visuals"].extend(extras)
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(formatted, load(REPORT), Path(temp_dir))
            self.assertEqual(result.status, "rendered", result.detail)
            asset_dir = Path(temp_dir) / "report_formatted_assets"
            self.assertTrue((asset_dir / "VIS-02.svg").is_file())
            self.assertTrue((asset_dir / "VIS-02.png").is_file())

    def test_format_runtime_enriches_chart_from_evidence_asset_ref(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            stage_dir = temp / "stage_4_format"
            handoff_dir = stage_dir / "handoff"
            handoff_dir.mkdir(parents=True)
            report = load(REPORT)
            report["delivery_target"] = "document"
            report["evidence_assets"] = [
                {
                    "evidence_id": "E1",
                    "asset_id": "T1-usage",
                    "ref": "E1:T1-usage",
                    "chart_ready": True,
                    "chart_data": {
                        "chart_type": "line",
                        "categories": ["2026-01-01", "2026-01-02"],
                        "series": [{"name": "DeepSeek", "values": [8, 16]}],
                    },
                }
            ]
            input_path = stage_dir / "input.json"
            write_json(input_path, report)
            write_json(
                stage_dir / "run_state.json",
                {
                    "run_id": "format-test",
                    "contract_profile": "v0_3",
                    "agent_id": "format",
                    "agent_name": "可视化",
                    "stage": 4,
                    "status": "init",
                    "current_step": "awaiting_gen_output",
                    "round_index": 0,
                    "input_path": str(input_path),
                    "output_dir": str(stage_dir),
                    "p0_open": [],
                    "p1_open": [],
                    "produced_artifacts": [],
                    "history": [],
                    "review_subagents_enabled": False,
                },
            )
            write_json(
                handoff_dir / "output_gen.json",
                {
                    "visuals": [
                        {
                            "visual_evidence_id": "VE-RUNTIME-01",
                            "section_heading": "一、成果保存与回访共同出现，但因果仍待验证",
                            "type": "chart",
                            "title": "DeepSeek 使用时长趋势",
                            "source_refs": ["E1"],
                            "required": False,
                            "placement": "section",
                        }
                    ]
                },
            )
            runner = StepRunner(ROOT, stage_dir, data_root=temp / "data", contract_profile="v0_3")
            result = RenderResult(
                status="rendered",
                fmt="document",
                fidelity="formatted",
                output_path=str(stage_dir / "report_formatted.docx"),
                file_bytes=1,
                unit_count=2,
                detail="preset=standard_business_brief",
            )

            with patch("presentation_agent.renderers.render_material", return_value=result):
                runner.commit()

            artifact = read_json(stage_dir / "artifact.json")
            visual = artifact["visuals"][0]
            self.assertEqual(visual["data"]["series"][0]["values"], [8, 16])
            self.assertIn("E1:T1-usage", visual["source_refs"])
            self.assertEqual(
                artifact["evidence_asset_enrichment"][0]["ref"],
                "E1:T1-usage",
            )
            self.assertEqual(
                artifact["render_result"]["detail"],
                "preset=standard_business_brief",
            )

    def test_line_chart_data_renders_as_document_visual(self) -> None:
        formatted = load(FORMATTED)
        formatted["visuals"][0]["type"] = "chart"
        formatted["visuals"][0]["data"] = {
            "chart_type": "line",
            "categories": ["2026-01-01", "2026-01-02", "2026-01-03"],
            "series": [
                {"name": "DeepSeek", "values": [8, 12, 16]},
                {"name": "豆包", "values": [9, 10, 11]},
            ],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(formatted, load(REPORT), Path(temp_dir))
            self.assertEqual(result.status, "rendered", result.detail)
            asset_dir = Path(temp_dir) / "report_formatted_assets"
            self.assertTrue((asset_dir / "VE-01.png").is_file())

    def test_real_case_chart_shapes_and_matrix_render_as_images(self) -> None:
        formatted = load(FORMATTED)
        formatted["visuals"] = [
            {
                "section_heading": "一、成果保存与回访共同出现，但因果仍待验证",
                "type": "chart",
                "title": "因子拆解",
                "source_refs": ["E-Q-01"],
                "data": {
                    "metrics": ["人均使用时长", "人均使用次数", "人均单次使用时长"],
                    "baseline_202505": [8, 4.65, 1.75],
                    "observation_202605_value_range": ["15-25+", "7-8", "2.1+"],
                    "change_range": ["~+100%", "+50-72%", "+20%"],
                },
            },
            {
                "section_heading": "二、可复用成果提供了值得优先验证的回访理由",
                "type": "chart",
                "title": "结构变化",
                "source_refs": ["E-Q-01"],
                "data": {
                    "categories": ["30分钟以上", "10-30分钟", "低频用户"],
                    "period_start_pct": [8.5, 15.0, 51.0],
                    "period_end_pct": [12.5, 17.6, 42.0],
                    "change_pp": ["+4", "+2.6", "-9"],
                },
            },
            {
                "section_heading": "二、可复用成果提供了值得优先验证的回访理由",
                "type": "matrix",
                "title": "四重结构性局限",
                "source_refs": ["E-Q-01"],
                "data": {
                    "dimensions": ["有效vs低效使用", "产出质量", "产品形态内生差异", "DAU稀释效应"],
                    "limitations": ["不能区分等待与有效交互", "不能衡量成功率", "跨品类基准有偏", "拉新改变分母"],
                },
            },
        ]
        with tempfile.TemporaryDirectory() as temp_dir:
            result = render_formatted_document_v2(formatted, load(REPORT), Path(temp_dir))
            self.assertEqual(result.status, "rendered", result.detail)
            asset_dir = Path(temp_dir) / "report_formatted_assets"
            for asset_id in ("VIS-01", "VIS-02", "VIS-03"):
                self.assertTrue((asset_dir / f"{asset_id}.png").is_file())
            from docx import Document

            self.assertEqual(len(Document(result.output_path).inline_shapes), 3)


if __name__ == "__main__":
    unittest.main()
